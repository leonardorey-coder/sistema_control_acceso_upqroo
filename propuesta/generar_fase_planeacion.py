from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).with_name("fase_planeacion_control_acceso_v2.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def paragraph_border_bottom(paragraph, color="2E74B5", size="8", space="4") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Fase de planeacion - Control Acceso v2"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Pagina ")
    add_field(footer, "PAGE")
    footer.add_run(" de ")
    add_field(footer, "NUMPAGES")
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


def add_title(doc: Document, text: str, size=24, after=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_BLUE


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr[idx].text = h
        set_cell_shading(hdr[idx], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_widths(table, widths)
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_BLUE
    r2 = p.add_run(text)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    set_table_widths(table, [6.5])
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    style_document(doc)
    set_header_footer(doc)

    add_title(doc, "UNIVERSIDAD POLITECNICA DE QUINTANA ROO", size=16, after=18)
    add_title(doc, "Administracion de Proyectos de Software", size=15, after=12)
    add_title(doc, "Fase de Planeacion del Proyecto", size=22, after=4)
    add_title(doc, "Sistema de Control de Acceso UPQROO v2", size=20, after=24)
    for label, value in [
        ("Carrera", "Ingenieria en Software / Tecnologias de la Informacion"),
        ("Materia", "Administracion de Proyectos de Software"),
        ("Docente", "Ing. Mejia Guzman Juan Carlos"),
        ("Equipo", "Equipo Control Acceso v2"),
        ("Grupo", "Por completar"),
        ("Cuatrimestre", "Por completar"),
        ("Fecha de entrega", "Julio de 2026"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(value)
    doc.add_page_break()

    doc.add_heading("Indice", level=1)
    add_caption(doc, "Tabla de contenido")
    add_table(doc, ["Seccion", "Pagina"], [
        ["1. Integrantes y roles", "4"],
        ["2. Descripcion general del proyecto", "4"],
        ["3. Planteamiento del problema", "5"],
        ["4. Objetivos", "6"],
        ["5. Antecedentes", "6"],
        ["6. Estado del arte", "7"],
        ["7. Justificacion", "8"],
        ["8. Variables del proyecto", "9"],
        ["9. Delimitacion del proyecto", "9"],
        ["10. Alcances", "10"],
        ["11. Limitaciones", "11"],
        ["12. Planeacion inicial", "11"],
        ["13. Fase de diseno", "13"],
        ["14. Planeacion economica y estimacion de costos", "17"],
        ["15. Factibilidad tecnica", "18"],
        ["16. Factibilidad operativa", "19"],
        ["17. Recursos del proyecto", "19"],
        ["18. Riesgos iniciales", "20"],
        ["19. Direccion, comunicacion y control del proyecto", "20"],
        ["20. Anexos", "21"],
        ["21. Conclusion", "21"],
        ["22. Bibliografia", "22"],
    ], [5.7, 0.8])
    doc.add_page_break()

    doc.add_heading("Indice de tablas y figuras", level=1)
    add_bullets(doc, [
        "Tabla 1. Integrantes y roles del proyecto",
        "Tabla 2. Comparativo de soluciones existentes",
        "Tabla 3. Variables principales del sistema",
        "Tabla 4. Delimitacion, alcances y exclusiones",
        "Tabla 5. Cronograma inicial de ocho semanas",
        "Tabla 6. Dinamica de asesoria semanal",
        "Tabla 7. Bitacora de seguimiento",
        "Tabla 8. Matriz de decisiones de la fase de diseno",
        "Tabla 9. Criterios de aceptacion del diseno",
        "Tabla 10. Presupuesto general estimado",
        "Tabla 11. Recursos gratuitos y recursos de pago",
        "Tabla 12. Recursos tecnologicos del proyecto",
        "Tabla 13. Riesgos iniciales",
        "Figura 1. Diseno como puente hacia desarrollo",
        "Figura 2. Flujo general del acceso auditable",
    ])
    doc.add_page_break()

    doc.add_heading("1. Integrantes y roles", level=1)
    add_body(doc, "La administracion de recursos humanos de la Unidad 2 establece que las personas son el recurso principal del proyecto. Por ello, el equipo se organiza por responsabilidades claras, comunicacion semanal y control de avances. Los nombres reales y fotografias de cada integrante deberan sustituir los campos de la tabla antes de la entrega final si el docente los solicita.")
    add_caption(doc, "Tabla 1. Integrantes y roles del proyecto")
    add_table(doc, ["Integrante", "Rol asignado", "Responsabilidades", "Fotografia"], [
        ["Integrante 1", "Lider del proyecto", "Planificar, coordinar reuniones, controlar alcance, validar entregables y resolver bloqueos.", "Espacio para fotografia"],
        ["Integrante 2", "Analista de requerimientos", "Levantar procesos de acceso, documentar usuarios, objetivos, variables, alcances y criterios de aceptacion.", "Espacio para fotografia"],
        ["Integrante 3", "Arquitecto backend y base de datos", "Definir API, seguridad, PostgreSQL, Drizzle, SQL atomico, auditoria e integridad.", "Espacio para fotografia"],
        ["Integrante 4", "Desarrollador frontend", "Construir panel administrativo, scanner, portal de usuario, formularios, tablas y experiencia movil.", "Espacio para fotografia"],
        ["Integrante 5", "QA y documentacion", "Preparar pruebas, revisar formato APS, mantener bitacora, evidencias, riesgos y bibliografia.", "Espacio para fotografia"],
    ], [1.0, 1.5, 3.2, 0.8])

    doc.add_heading("2. Descripcion general del proyecto", level=1)
    add_body(doc, "Control Acceso v2 es un sistema web institucional para administrar entradas, salidas, asistencia, credenciales QR, visitantes, vehiculos, permisos, auditoria y reportes operativos dentro de la Universidad Politecnica de Quintana Roo. El sistema se dirige a guardias, administradores, super administradores, alumnos, docentes, personal administrativo, visitantes y responsables de analitica institucional.")
    add_body(doc, "El proyecto es un sistema de informacion administrativo y operativo. Su proposito principal es reducir el uso de registros manuales, mejorar la seguridad del campus, conservar evidencia auditable de cada acceso y ofrecer una experiencia rapida para el punto fisico de entrada.")
    add_callout(doc, "Objetivo administrativo", "Presentar una planeacion viable antes de continuar el desarrollo, integrando antecedentes, estado del arte, delimitacion, cronograma, presupuesto, factibilidad, recursos y riesgos.")

    doc.add_heading("3. Planteamiento del problema", level=1)
    add_body(doc, "La situacion actual de control de acceso en instituciones educativas suele depender de listas, registros manuales, credenciales visuales o validaciones aisladas. Ese modelo dificulta conocer en tiempo real quien entro, quien salio, que vehiculo tiene permiso, que visitante fue autorizado, que usuario presento un QR vigente y que operador realizo la validacion.")
    add_body(doc, "El problema central es la falta de una trazabilidad unificada, segura y consultable para accesos peatonales, vehiculares y temporales. Cuando la evidencia se distribuye en hojas, mensajes, capturas o sistemas inconexos, aparecen errores de captura, duplicidad, reingresos no controlados, poca auditoria y baja capacidad de reaccion ante incidentes.")
    add_body(doc, "La necesidad informatica consiste en consolidar un sistema con autenticacion segura, QR firmado de corta vida, scanner operativo, registros con entrada y salida, cierre automatico, auditoria administrativa, portal de usuario, administracion de vehiculos y reportes filtrables.")

    doc.add_heading("4. Objetivos", level=1)
    doc.add_heading("4.1 Objetivo general", level=2)
    add_body(doc, "Planear y desarrollar una plataforma web institucional de control de acceso para la UPQROO que permita registrar, validar, auditar y consultar entradas y salidas de personas, visitantes y vehiculos mediante credenciales QR seguras, administracion centralizada y evidencia operativa confiable.")
    doc.add_heading("4.2 Objetivos especificos", level=2)
    add_numbered(doc, [
        "Definir el alcance funcional de Control Acceso v2 considerando scanner, panel administrativo, portal de usuario, Hot-QR, vehiculos, asistencia, auditoria y reportes.",
        "Diseñar una arquitectura tecnica viable basada en SvelteKit, TypeScript, Bun, Hono, PostgreSQL, Drizzle ORM y SQL manual versionado para reglas atomicas.",
        "Establecer un cronograma inicial de ocho semanas con actividades, responsables, duracion y criterios de seguimiento.",
        "Estimar costos profesionales de desarrollo, infraestructura, dispositivos, almacenamiento, capacitacion y contingencia.",
        "Identificar riesgos iniciales y estrategias de mitigacion relacionadas con seguridad, adopcion operativa, datos, hardware, tiempos e infraestructura.",
    ])

    doc.add_heading("5. Antecedentes", level=1)
    add_body(doc, "La Unidad 1 define los antecedentes como la informacion previa que explica como surgio el problema, que se ha hecho antes y cual es la situacion actual. En este proyecto, el antecedente principal es la existencia de una version legacy funcional que resolvia flujos basicos de administracion y scanner, pero con limitaciones de seguridad, mantenibilidad y separacion de responsabilidades.")
    add_body(doc, "El repo actual de Control Acceso v2 ya documenta una evolucion hacia un monorepo con API Bun + Hono + TypeScript, frontend SvelteKit, PostgreSQL con Drizzle y SQL manual versionado. Tambien conserva requisitos operativos del legacy: header institucional, scanner separado, resultado visual claro, foto, datos de persona, tablas densas, formularios compactos, Hot-QR y administracion de personas.")
    add_body(doc, "La planeacion parte de conservar lo que el usuario operativo ya reconoce, pero corrigiendo debilidades: tokens en cliente, QR con informacion sensible, falta de auditoria granular, poca modularidad, ausencia de pruebas por modulo y dificultad para extender vehiculos, portal, asistencia y reportes.")

    doc.add_heading("6. Estado del arte", level=1)
    add_body(doc, "El estado del arte permite comparar soluciones existentes, analizar ventajas y detectar oportunidades. Para Control Acceso v2 se revisaron soluciones comerciales de control de acceso fisico, credenciales moviles y administracion en la nube. La comparacion no busca copiarlas, sino identificar buenas practicas aplicables a un contexto universitario con presupuesto controlado.")
    add_caption(doc, "Tabla 2. Comparativo de soluciones existentes")
    add_table(doc, ["Solucion", "Funciones principales", "Ventajas", "Desventajas", "Aporte para v2"], [
        ["Kisi", "Control de acceso en la nube, app movil, hardware y software unificados, integraciones, reportes y administracion remota.", "Operacion centralizada y escalable; buena referencia para permisos, usuarios y ubicaciones.", "Dependencia de hardware/servicio externo; costo recurrente; menor personalizacion institucional.", "Control Acceso v2 debe ofrecer administracion remota y permisos claros, pero con ownership de datos institucionales."],
        ["HID", "Credenciales fisicas y moviles, tarjetas inteligentes, biometria, identidad, visitor management y soluciones para universidades.", "Solidez industrial, compatibilidad con credenciales seguras y experiencia en campus.", "Implementacion puede requerir hardware especializado, integradores y costos altos.", "v2 toma la idea de identidad confiable y credencial segura, pero usa QR firmado y portal web como alternativa academica viable."],
        ["Verkada", "Control de acceso hibrido en nube, video, integraciones, credenciales moviles, alertas y administracion de dispositivos.", "Vision integral de seguridad fisica, eventos y contexto visual.", "Costo enterprise, dependencia de plataforma cerrada y sensibilidad de datos de vigilancia.", "v2 debe integrar auditoria, eventos y contexto visual sin exponer mas datos personales de los necesarios."],
        ["Avigilon / Openpath", "Credenciales moviles, tarjetas, PIN, pases de invitado, software de gestion, integracion con video y componentes fisicos.", "Muestra componentes clave: lectores, credenciales, software y hardware de apertura.", "Orientado a infraestructura fisica profesional; costos y despliegue superan una primera fase academica.", "v2 puede preparar el modelo para gates fisicos futuros sin exigirlos en el alcance inicial."],
    ], [1.0, 1.6, 1.45, 1.45, 1.0])
    add_body(doc, "Analisis comparativo: las soluciones comerciales priorizan hardware, credenciales moviles, administracion remota, integraciones y escalabilidad. Control Acceso v2 aportara una solucion ajustada a UPQROO, con QR dinamico firmado, portal de usuario, auditoria, reglas atomicas en PostgreSQL, lenguaje visual institucional, Hot-QR y vehiculos, manteniendo costos iniciales menores que una plataforma enterprise.")

    doc.add_heading("7. Justificacion", level=1)
    add_body(doc, "El proyecto es importante porque el acceso fisico a un campus es un proceso operativo y de seguridad. Un registro confiable permite saber quien se encuentra dentro, que entradas quedaron abiertas, que visitante fue autorizado, que vehiculo tiene permiso y que operador intervino. Esto impacta seguridad, administracion, reportes y toma de decisiones.")
    add_body(doc, "Los usuarios beneficiados son guardias, administradores, super administradores, comunidad universitaria, visitantes autorizados y direccion institucional. El beneficio principal es reducir friccion en el punto de acceso sin perder control: el guardia escanea, el sistema valida y registra, y la administracion consulta evidencia.")
    add_body(doc, "El desarrollo se justifica porque aprovecha tecnologia web accesible, evita depender desde el inicio de hardware propietario y permite adaptar el sistema a reglas locales de la universidad, horarios, carreras, tipos de persona, permisos vehiculares y excepciones.")

    doc.add_heading("8. Variables del proyecto", level=1)
    add_caption(doc, "Tabla 3. Variables principales del sistema")
    add_table(doc, ["Variable", "Descripcion", "Tipo"], [
        ["Persona", "Alumno, docente, administrativo, visitante o perfil autorizado.", "Entidad principal"],
        ["Matricula/identificador", "Dato visible para busqueda y operacion; no debe funcionar como secreto.", "Texto"],
        ["Tipo de persona", "Categoria con reglas: carrera requerida, portal, asistencia, permiso vehicular.", "Catalogo"],
        ["Credencial QR", "Token opaco o QR firmado de corta vida para validar acceso.", "Seguridad"],
        ["JTI", "Identificador unico de token firmado para evitar replay.", "UUID/texto"],
        ["Registro de acceso", "Entrada, salida, estado, modo, scanner, hora y relacion con persona/vehiculo.", "Transaccional"],
        ["Evento de scanner", "Evidencia tecnica de aceptacion o rechazo.", "Auditoria"],
        ["Vehiculo", "Placa, propietario, estado, tipo y permiso vinculado.", "Entidad"],
        ["Hot-QR", "Credencial temporal de visitante con motivo, vigencia y usos.", "Temporal"],
        ["Administrador", "Usuario autorizado para operar panel, scanner o configuracion.", "Seguridad"],
        ["Asistencia", "Registro potencial o confirmado asociado a horarios y materias.", "Academico-operativo"],
        ["Configuracion", "Zona horaria, QR firmado, expiraciones, worker y reglas del sistema.", "Parametro"],
    ], [1.6, 3.6, 1.3])

    doc.add_heading("9. Delimitacion del proyecto", level=1)
    add_caption(doc, "Tabla 4. Delimitacion, alcances y exclusiones")
    add_table(doc, ["Si desarrollara", "Quedara fuera de la fase inicial"], [
        ["Panel administrativo con personas, registros, asistencia, Hot-QR, vehiculos, administradores y configuracion.", "Integracion directa con torniquetes o cerraduras fisicas reales."],
        ["Scanner web con lectura de QR firmado, token opaco y captura manual autorizada.", "Aplicacion movil nativa para iOS/Android."],
        ["Portal de usuario con QR dinamico, historial y dispositivos.", "Reconocimiento facial automatizado como mecanismo de autorizacion."],
        ["PostgreSQL, Drizzle, SQL atomico, auditoria, sesiones y cookies httpOnly.", "Migracion historica completa de todos los datos legacy si no hay dataset validado."],
        ["Reportes y filtros operativos basicos para acceso, rechazo, asistencia y vehiculos.", "BI avanzado o tableros directivos con prediccion estadistica."],
    ], [3.25, 3.25])

    doc.add_heading("10. Alcances", level=1)
    add_bullets(doc, [
        "Autenticacion administrativa y de portal con sesiones seguras y cookies httpOnly.",
        "Administracion de personas, tipos de persona, carreras, fotos y estados.",
        "Generacion, rotacion y revocacion de credenciales QR personales, temporales y vehiculares.",
        "Scanner para entrada, salida, rechazo, QR temporal, Hot-QR, vehiculo y captura manual auditada.",
        "Registro atomico en PostgreSQL de accesos, eventos, asistencia e integridad de cadena hash.",
        "Gestion de vehiculos, permisos y QR vehicular firmado u opaco.",
        "Hot-QR para visitantes con vigencia, uso maximo, revocacion y motivo.",
        "Panel administrativo con tablas, filtros, paginacion, badges y acciones auditadas.",
        "Worker operativo para expiraciones, cierres automaticos, limpieza y mantenimiento.",
        "Pruebas con bun:test y verificacion funcional de flujos criticos.",
    ])

    doc.add_heading("11. Limitaciones", level=1)
    add_bullets(doc, [
        "Tiempo academico limitado a ocho semanas de planeacion, desarrollo y cierre documental.",
        "Presupuesto controlado; se privilegia software libre y servicios de bajo costo.",
        "La fase inicial depende de equipos de computo o tablets disponibles para scanner.",
        "La precision de reportes historicos depende de la calidad de datos capturados o migrados.",
        "El despliegue productivo requiere disponibilidad estable de internet, servidor y base de datos.",
        "La adopcion operativa requiere capacitacion de guardias y administradores.",
        "No se cubriran integraciones fisicas con barreras automatizadas hasta una fase posterior.",
    ])

    doc.add_heading("12. Planeacion inicial", level=1)
    add_body(doc, "La planeacion inicial se construye con base en las Unidades 1 y 2: objetivos, actividades, recursos, tiempo, responsables, comunicacion y control. Se propone un ciclo de ocho semanas compatible con la propuesta de asignacion de asesorias, donde cada reunion revisa avances, nuevo trabajo, retroalimentacion y tareas siguientes.")
    add_caption(doc, "Tabla 5. Cronograma inicial de ocho semanas")
    add_table(doc, ["Semana", "Actividad", "Responsable", "Duracion", "Entregable"], [
        ["1", "Presentacion del proyecto, problema, actores y alcance inicial.", "Lider + Analista", "1 semana", "Propuesta inicial validada"],
        ["2", "Objetivos, antecedentes, estado del arte y justificacion.", "Analista + QA", "1 semana", "Secciones teoricas corregidas"],
        ["3", "Variables, delimitacion, alcances y limitaciones.", "Analista + Arquitecto", "1 semana", "Modelo de variables y alcance"],
        ["4", "Planeacion inicial y planeacion economica.", "Lider + QA", "1 semana", "Cronograma y presupuesto"],
        ["5", "Factibilidad tecnica y operativa.", "Arquitecto + Frontend", "1 semana", "Factibilidades documentadas"],
        ["6", "Recursos del proyecto, riesgos y estrategia PDCA.", "Todo el equipo", "1 semana", "Matriz de recursos y riesgos"],
        ["7", "Integracion del documento, revision de formato APS y evidencias.", "QA + Lider", "1 semana", "Documento completo revisado"],
        ["8", "Presentacion final, ajustes y entrega.", "Todo el equipo", "1 semana", "Documento final y exposicion"],
    ], [0.55, 2.25, 1.25, 0.8, 1.65])
    add_caption(doc, "Tabla 6. Dinamica de asesoria semanal")
    add_table(doc, ["Actividad", "Tiempo", "Proposito"], [
        ["Revision de avances de la semana anterior", "10 min", "Comparar lo planeado contra lo realizado."],
        ["Presentacion del nuevo avance", "20 min", "Mostrar entregables, decisiones y bloqueos."],
        ["Retroalimentacion y correcciones", "20 min", "Ajustar alcance, formato, contenido y riesgos."],
        ["Definicion de tareas siguientes", "10 min", "Asignar responsables y compromisos."],
    ], [3.0, 1.0, 2.5])
    add_caption(doc, "Tabla 7. Bitacora de seguimiento")
    add_table(doc, ["Semana", "Avance esperado", "Cumplio", "Observaciones", "Firma"], [
        ["1", "Idea del proyecto y problema", "", "", ""],
        ["2", "Objetivos, antecedentes y estado del arte", "", "", ""],
        ["3", "Variables, delimitacion y alcances", "", "", ""],
        ["4", "Planeacion y costos", "", "", ""],
        ["5", "Factibilidad tecnica y operativa", "", "", ""],
        ["6", "Recursos y riesgos", "", "", ""],
        ["7", "Documento completo", "", "", ""],
        ["8", "Presentacion final", "", "", ""],
    ], [0.7, 2.2, 0.75, 2.05, 0.8])

    doc.add_heading("13. Fase de diseno", level=1)
    add_body(doc, "La fase de diseno funciona como puente entre la planeacion y el desarrollo. Antes de programar nuevas pantallas o endpoints, el equipo debe convertir los objetivos, alcances, restricciones y riesgos en decisiones concretas sobre plataforma, usuarios, tiempo, datos/latencia y presupuesto. Esta fase evita que el desarrollo avance con supuestos incompletos.")
    add_caption(doc, "Figura 1. Diseno como puente hacia desarrollo")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("Plataforma + Usuarios + Tiempo + Datos/Latencia + Presupuesto -> Diseno -> Desarrollo")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    add_callout(doc, "Decision de diseno", "Control Acceso v2 no debe pasar a desarrollo solo por tener una lista de funciones. Cada funcion debe tener usuario objetivo, plataforma de uso, restriccion de tiempo, impacto en datos/latencia y costo estimado.")
    add_caption(doc, "Tabla 8. Matriz de decisiones de la fase de diseno")
    add_table(doc, ["Tema", "Decision para Control Acceso v2", "Criterio de diseno", "Evidencia esperada"], [
        ["Plataforma", "Aplicacion web SvelteKit para panel, portal y scanner; API Bun/Hono; PostgreSQL como fuente de verdad; storage local/S3/R2 para fotos.", "El sistema debe operar en navegador moderno, tablets o equipos de escritorio sin exigir app movil nativa en fase inicial.", "Mapa de rutas, modulos, endpoints /api/v1, variables de entorno y diagrama de despliegue."],
        ["Usuarios", "Separar guardia/scanner, administrador operativo, super administrador, usuario de portal, visitante y analitica institucional.", "Cada pantalla debe mostrar solo acciones y datos necesarios para su rol; el scanner prioriza rapidez y lectura visual.", "Historias de usuario, matriz de permisos y flujos por rol."],
        ["Tiempo", "Diseñar para escaneo rapido, respuesta inmediata y ciclo de desarrollo de ocho semanas con entregables verificables.", "El flujo de scanner debe minimizar pasos: leer QR, validar, mostrar aceptado/rechazado y permitir continuar.", "Cronograma, criterios de avance semanal y pruebas de tiempo de respuesta."],
        ["Datos / Latencia", "Mantener decisiones atomicas en PostgreSQL: entrada/salida, consumo de jti, asistencia, eventos y auditoria.", "La latencia del scan depende de consultas, indices, pool y worker; no se deben guardar tokens completos ni datos sensibles innecesarios.", "Modelo de datos, indices, pruebas de performance y reglas de privacidad."],
        ["Presupuesto", "Priorizar software libre, servicios administrados de bajo costo y dispositivos scanner reutilizables.", "El diseno debe sostenerse con el presupuesto estimado; funciones de alto costo pasan a fases posteriores.", "Presupuesto, analisis de reduccion del 20 % y lista de funciones postergables."],
    ], [1.05, 2.1, 2.0, 1.35])
    doc.add_heading("13.1 Diseno de plataforma", level=2)
    add_body(doc, "La plataforma se disena como un monorepo web porque el proyecto requiere panel administrativo, scanner y portal en una misma base tecnica. SvelteKit permite rutas y componentes reutilizables; Bun y Hono permiten API REST ligera; PostgreSQL conserva integridad, auditoria y transacciones. La decision evita duplicar una app movil nativa durante la fase inicial y concentra el esfuerzo en un flujo web usable.")
    add_bullets(doc, [
        "Panel administrativo: gestion de personas, QR, registros, asistencia, Hot-QR, vehiculos, administradores y configuracion.",
        "Scanner: vista operativa separada, optimizada para camara, resultado visual y continuidad de escaneo.",
        "Portal: QR dinamico, historial propio, dispositivos y cambio de password.",
        "API: contratos versionados /api/v1, validacion con Zod, sesiones httpOnly y auditoria.",
        "Base de datos: Drizzle para schema tipado y SQL manual versionado para reglas atomicas.",
    ])
    doc.add_heading("13.2 Diseno centrado en usuarios", level=2)
    add_body(doc, "El diseno debe partir de los usuarios reales. El guardia necesita velocidad y claridad, no una pantalla cargada de administracion. El administrador necesita busqueda, filtros y acciones auditadas. El usuario de portal necesita consultar su QR sin tocar configuraciones. El super administrador necesita control de roles, sesiones y auditoria.")
    add_bullets(doc, [
        "Guardia: camara visible, estado de espera, resultado aceptado/rechazado, foto y motivo claro.",
        "Administrador operativo: formularios compactos, tablas densas, busqueda por matricula, nombre, placa y estado.",
        "Super administrador: proteccion de ultimo superadmin, sesiones, auditoria y acciones sensibles con motivo.",
        "Usuario de portal: acceso movil simple, QR vigente, historial y dispositivos vinculados.",
        "Visitante: QR temporal con vigencia, usos, motivo y revocacion.",
    ])
    doc.add_heading("13.3 Diseno por tiempo", level=2)
    add_body(doc, "El tiempo se considera en dos niveles: tiempo de proyecto y tiempo de operacion. En proyecto, cada semana debe cerrar evidencia concreta. En operacion, el scanner debe responder rapido porque el punto de acceso no puede convertirse en cuello de botella.")
    add_bullets(doc, [
        "Tiempo de proyecto: mantener entregables por semana, revision de avance y correccion temprana.",
        "Tiempo de scanner: evitar pantallas intermedias, confirmar resultado en una sola vista y permitir reintento rapido.",
        "Tiempo de mantenimiento: worker para expiraciones, salidas automaticas y limpieza sin intervencion manual.",
        "Tiempo de soporte: logs, auditoria y reportes deben permitir diagnosticar errores sin revisar manualmente la base de datos.",
    ])
    doc.add_heading("13.4 Diseno de datos y latencia", level=2)
    add_body(doc, "El diseno de datos debe proteger la seguridad y sostener el rendimiento. En Control Acceso v2, la matricula no debe funcionar como secreto; los QR firmados deben tener expiracion corta; el jti se consume atomicamente para evitar replay; y los eventos deben registrar evidencia sin persistir tokens completos.")
    add_bullets(doc, [
        "Datos minimos: mostrar al guardia solo identidad necesaria, foto, estado, tipo, carrera/rol y razon de rechazo.",
        "Latencia: optimizar consultas de scan, registros del dia, busqueda de personas, busqueda vehicular y asistencia.",
        "Integridad: mantener cadena hash y auditoria para accesos aceptados y rechazados.",
        "Privacidad: usar storage protegido, URL firmada temporal y permisos por sesion.",
        "Escalabilidad: controlar pool PostgreSQL, indices y coalescencia de eventos WebSocket.",
    ])
    doc.add_heading("13.5 Diseno condicionado por presupuesto", level=2)
    add_body(doc, "El presupuesto condiciona el diseno porque no todas las funciones pueden entrar en la primera fase. La estrategia es proteger el nucleo: autenticacion, personas, QR seguro, scanner, registros, auditoria, Hot-QR, vehiculos basicos y reportes operativos. Las integraciones fisicas, BI avanzado, app movil nativa y automatizaciones costosas se dejan para una etapa posterior.")
    add_caption(doc, "Tabla 9. Criterios de aceptacion del diseno")
    add_table(doc, ["Dimension", "Criterio minimo antes de desarrollo", "Riesgo si se omite"], [
        ["Plataforma", "Rutas, modulos, API, base de datos y despliegue definidos.", "Reescrituras y dependencias improvisadas."],
        ["Usuarios", "Historias por rol y permisos principales documentados.", "Pantallas genericas que no resuelven el flujo operativo."],
        ["Tiempo", "Cronograma y tiempos criticos de scanner identificados.", "Retrasos y experiencia lenta en punto de acceso."],
        ["Datos / Latencia", "Modelo, indices, eventos, privacidad y jti definidos.", "Rechazos erroneos, replay, consultas lentas o filtracion de datos."],
        ["Presupuesto", "Costo base, recorte del 20 % y funciones postergables claros.", "Sobrecosto o promesas fuera del alcance academico."],
    ], [1.25, 3.0, 2.25])

    doc.add_heading("14. Planeacion economica y estimacion de costos", level=1)
    add_body(doc, "El presupuesto se presenta en un escenario profesional aunque el proyecto sea academico. Los costos son aproximados en MXN y priorizan herramientas gratuitas cuando no comprometen seguridad, disponibilidad ni mantenimiento.")
    add_caption(doc, "Tabla 10. Presupuesto general estimado")
    add_table(doc, ["Concepto", "Cantidad", "Costo unitario", "Costo total", "Justificacion"], [
        ["Horas de desarrollo", "520 h", "$180", "$93,600", "Analisis, backend, frontend, base de datos, pruebas y documentacion."],
        ["Servidor/VPS o hosting app", "12 meses", "$800", "$9,600", "Ejecucion de API, frontend y worker operativo."],
        ["PostgreSQL administrado", "12 meses", "$600", "$7,200", "Base de datos relacional, backups y disponibilidad."],
        ["Dominio", "1", "$800", "$800", "Acceso institucional facil de recordar."],
        ["Storage S3/R2", "12 meses", "$150", "$1,800", "Fotos de perfil, evidencias y archivos protegidos."],
        ["Tablets o celulares scanner", "2", "$4,500", "$9,000", "Dispositivos para puntos de acceso."],
        ["Credenciales/impresion inicial", "1 lote", "$1,000", "$1,000", "Material de prueba y operacion inicial."],
        ["Capacitacion", "10 h", "$180", "$1,800", "Sesion para guardias, administradores y soporte."],
        ["Contingencia 10%", "1", "$12,480", "$12,480", "Riesgos de cambios, ajustes o servicios adicionales."],
        ["Costo total estimado", "", "", "$137,280", "Total profesional aproximado."],
    ], [1.35, 0.7, 0.85, 0.85, 2.75])
    add_body(doc, "El recurso mas costoso es el trabajo de desarrollo, porque concentra analisis, codificacion, pruebas, correccion de defectos, documentacion y administracion del proyecto. Si el presupuesto se redujera 20 %, se mantendrian backend, base de datos, scanner y auditoria; se postergarian reportes avanzados, model-viewer vehicular, personalizacion visual secundaria y compra de dispositivos adicionales.")
    add_caption(doc, "Tabla 11. Recursos gratuitos y recursos de pago")
    add_table(doc, ["Recursos gratuitos", "Recursos de pago"], [
        ["Visual Studio Code, Git, GitHub, Bun, TypeScript, SvelteKit, Hono, Drizzle ORM, Zod, Tailwind CSS, html5-qrcode, qrcode, jose.", "Servidor o VPS, dominio, PostgreSQL administrado, almacenamiento R2/S3, dispositivos scanner, capacitacion y contingencia."],
    ], [3.25, 3.25])

    doc.add_heading("15. Factibilidad tecnica", level=1)
    add_body(doc, "La factibilidad tecnica es alta porque el proyecto ya cuenta con stack definido y componentes implementados en el repositorio: monorepo, API Bun/Hono, frontend SvelteKit, PostgreSQL, Drizzle, SQL atomico, QR firmado, auditoria y worker. Las herramientas son actuales, ampliamente documentadas y compatibles con despliegues de bajo costo.")
    add_bullets(doc, [
        "Hardware requerido: laptops de desarrollo, servidor/VPS, dispositivos moviles o tablets para scanner, camara y conectividad.",
        "Software requerido: Bun, TypeScript, SvelteKit, Hono, PostgreSQL, Drizzle Kit, Git, navegador moderno y servicios de storage.",
        "Infraestructura: base de datos con backups, dominio, HTTPS, variables de entorno seguras y proceso worker permanente.",
        "Conocimientos: desarrollo web, API REST, seguridad de sesiones, SQL, modelado relacional, pruebas, despliegue y administracion basica de servidores.",
    ])

    doc.add_heading("16. Factibilidad operativa", level=1)
    add_body(doc, "La factibilidad operativa es favorable si la interfaz conserva el lenguaje legacy que los operadores ya reconocen: header institucional, tabs superiores, formularios compactos, scanner separado, respuesta visual inmediata y tablas densas. La adopcion dependera menos de tecnologia compleja y mas de capacitacion, claridad de roles y continuidad del flujo actual.")
    add_body(doc, "La capacitacion minima debe cubrir: inicio de sesion, uso de scanner, interpretacion de aceptado/rechazado, busqueda de registros, creacion de Hot-QR, gestion de personas, revision de vehiculos, cierre automatico, revocacion de credenciales y reporte de incidencias.")

    doc.add_heading("17. Recursos del proyecto", level=1)
    add_caption(doc, "Tabla 12. Recursos tecnologicos del proyecto")
    add_table(doc, ["Categoria", "Recursos", "Uso"], [
        ["Humanos", "Lider, analista, backend, frontend, QA/documentacion.", "Planeacion, desarrollo, validacion y entrega."],
        ["Frontend", "SvelteKit, TypeScript, Tailwind, Vite, html5-qrcode, qrcode.", "Panel admin, scanner y portal."],
        ["Backend", "Bun, Hono, Zod, jose, API REST /api/v1.", "Servicios, autenticacion, validacion y QR firmado."],
        ["Datos", "PostgreSQL, Drizzle ORM, Drizzle Kit, SQL manual versionado.", "Schema, migraciones, reglas atomicas, auditoria y reportes."],
        ["Seguridad", "Cookies httpOnly, Bun.password, sesiones hasheadas, rate limit, JWKS.", "Proteccion de usuarios, admins y credenciales."],
        ["Operacion", "Worker, storage local/S3/R2, logs, performance harness.", "Expiraciones, archivos, mantenimiento y medicion."],
    ], [1.2, 2.8, 2.5])

    doc.add_heading("18. Riesgos iniciales", level=1)
    add_caption(doc, "Tabla 13. Riesgos iniciales")
    add_table(doc, ["Riesgo", "Probabilidad", "Impacto", "Estrategia de mitigacion"], [
        ["QR reutilizado o replay de token", "Media", "Alto", "Usar QR firmado con expiracion corta, jti unico y consumo atomico en PostgreSQL."],
        ["Falla de internet en punto de acceso", "Media", "Alto", "Definir modo contingencia manual auditado y sincronizacion posterior si aplica."],
        ["Resistencia de guardias o administradores", "Media", "Medio", "Capacitacion corta, interfaz legacy-style y pruebas piloto con retroalimentacion."],
        ["Datos incompletos de personas o vehiculos", "Alta", "Medio", "Validacion CSV, campos obligatorios, estados claros y bitacora de correccion."],
        ["Sobrecarga de base de datos en horas pico", "Media", "Alto", "Indices, pool controlado, pruebas de performance y SQL atomico optimizado."],
        ["Perdida o exposicion de fotos/datos personales", "Baja", "Alto", "Storage protegido, URLs firmadas, permisos por sesion y minimo dato visible."],
        ["Cambios de alcance durante el curso", "Media", "Medio", "Control de cambios semanal y delimitacion documentada."],
        ["Worker detenido o mal configurado", "Media", "Alto", "Monitoreo, scripts de operacion y validacion de expiraciones/cierres automaticos."],
    ], [1.7, 0.9, 0.75, 3.15])

    doc.add_heading("19. Direccion, comunicacion y control del proyecto", level=1)
    add_body(doc, "La Unidad 2 define la direccion de proyectos como la aplicacion de conocimientos, habilidades y herramientas para planificar, ejecutar y controlar. En Control Acceso v2, el lider debe mantener decisiones informadas sobre tiempo, alcance, costos y calidad.")
    add_body(doc, "Se aplicara el ciclo PDCA: Planear las actividades semanales, Hacer los entregables, Verificar avances y pruebas, y Actuar corrigiendo desviaciones. Este enfoque reduce errores y permite mejora continua sin perder control del alcance.")
    add_bullets(doc, [
        "Comunicacion: reunion semanal de 60 minutos, bitacora, tablero de tareas y acuerdos por escrito.",
        "Gestion de conflictos: escuchar partes, comparar contra objetivos, negociar alcance y documentar decisiones.",
        "Seguimiento del tiempo: cronograma de ocho semanas y revision de lo planeado contra lo ejecutado.",
        "Seguimiento de costos: presupuesto base, horas estimadas, recursos gratuitos y escenarios de reduccion.",
    ])

    doc.add_heading("20. Anexos", level=1)
    add_caption(doc, "Figura 2. Flujo general del acceso auditable")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("Usuario presenta QR -> Scanner lee credencial -> API valida sesion y token -> PostgreSQL decide entrada/salida/rechazo -> Auditoria registra evento -> Panel/portal consultan evidencia")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    add_body(doc, "Anexo A. Fuentes del curso analizadas: Unidad 1 parte 2, Unidad 2, Guia para la documentacion APS y Propuesta de Asignacion de Equipos. Estas fuentes definieron el formato, estructura, secciones, cronograma de asesorias y criterios de planeacion.")
    add_body(doc, "Anexo B. Fuentes del proyecto analizadas: README, PRD, plan de migracion de stack, plan tecnologico, feature de entrada/salida auditable y schema de base de datos. Estas fuentes definieron el alcance tecnico real de Control Acceso v2.")

    doc.add_heading("21. Conclusion", level=1)
    add_body(doc, "Una administracion adecuada de proyectos de software permite convertir una necesidad operativa en un plan ejecutable. Para Control Acceso v2, la planeacion aclara problema, objetivos, actores, variables, alcance, limitaciones, fase de diseno, tiempo, costos, recursos, factibilidad y riesgos antes de continuar el desarrollo.")
    add_body(doc, "El principal reto es equilibrar seguridad, rapidez operativa, datos confiables, latencia baja y costo. El sistema debe ser suficientemente robusto para proteger credenciales, sesiones y auditoria, pero lo bastante simple para que guardias y administradores lo adopten sin friccion. Desde la factibilidad tecnica, operativa y economica, el proyecto es viable si se respeta el alcance inicial, se priorizan los flujos criticos y se deja la integracion fisica avanzada para fases posteriores.")

    doc.add_heading("22. Bibliografia", level=1)
    refs = [
        "[1] J. C. Mejia Guzman, Administracion de Proyectos de Software U1 Parte 2, Universidad Politecnica de Quintana Roo, 2026.",
        "[2] J. C. Mejia Guzman, Administracion de Proyectos de Software U2, Universidad Politecnica de Quintana Roo, 2026.",
        "[3] Universidad Politecnica de Quintana Roo, Guia para la administracion de proyectos de software, 2026.",
        "[4] J. C. Mejia Guzman, Propuesta de Asignacion de Equipos, 2026.",
        "[5] Sistema de Control de Acceso UPQROO v2, README.md y documentacion tecnica del repositorio, 2026.",
        "[6] Kisi, 'Cloud-Based Access Control & Security Platform.' Disponible: https://www.getkisi.com/. Consultado: julio de 2026.",
        "[7] HID Global, 'Powering Trusted Identities.' Disponible: https://www.hidglobal.com/. Consultado: julio de 2026.",
        "[8] Verkada, 'Access control for the modern enterprise.' Disponible: https://www.verkada.com/access-control/. Consultado: julio de 2026.",
        "[9] Avigilon, 'Access Control Systems & Solutions.' Disponible: https://www.avigilon.com/access-control. Consultado: julio de 2026.",
    ]
    for ref in refs:
        add_body(doc, ref)

    doc.core_properties.title = "Fase de Planeacion - Control Acceso v2"
    doc.core_properties.subject = "Administracion de Proyectos de Software"
    doc.core_properties.author = "Equipo Control Acceso v2"
    doc.core_properties.keywords = "UPQROO, control de acceso, planeacion, APS"
    doc.save(OUT)


if __name__ == "__main__":
    build()
