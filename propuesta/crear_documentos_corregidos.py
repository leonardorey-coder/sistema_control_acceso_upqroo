from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from lxml import etree
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement as DocxElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "fase_planeacion_control_acceso_v2.docx"
CORRECTED = ROOT / "fase_planeacion_control_acceso_v2_corregida.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
W = f"{{{W_NS}}}"


OBJECTIVE_OLD = (
    "Planear y desarrollar una plataforma web institucional de control de acceso para la "
    "UPQROO que permita registrar, validar, auditar y consultar entradas y salidas de personas, "
    "visitantes y vehiculos mediante credenciales QR seguras, administracion centralizada y "
    "evidencia operativa confiable."
)
OBJECTIVE_NEW = (
    "Desarrollar una plataforma web institucional de control de acceso para la UPQROO que "
    "audite las entradas y salidas de personas, visitantes y vehiculos mediante credenciales "
    "QR seguras, administracion centralizada y evidencia operativa confiable."
)

OPERABILITY_OLD_1 = (
    "La factibilidad operativa es favorable si la interfaz conserva el lenguaje legacy que los "
    "operadores ya reconocen: header institucional, tabs superiores, formularios compactos, "
    "scanner separado, respuesta visual inmediata y tablas visuales no tan densas. La adopcion "
    "dependera menos de tecnologia compleja y mas de capacitacion, claridad de roles y continuidad "
    "del flujo actual."
)
OPERABILITY_NEW_1 = (
    "La factibilidad operativa es favorable si la interfaz conserva el lenguaje legacy que los "
    "operadores ya reconocen: header institucional, tabs superiores, formularios compactos, "
    "scanner separado, respuesta visual inmediata y tablas visuales no tan densas. La adopcion "
    "dependera principalmente de la capacitacion, la claridad de roles y la continuidad del flujo actual."
)
OPERABILITY_OLD_2 = (
    "La capacitacion minima debe cubrir: inicio de sesion, uso de scanner, interpretacion de "
    "aceptado/rechazado, busqueda de registros, creacion de Hot-QR, gestion de personas, revision "
    "de vehiculos, cierre automatico, revocacion de credenciales y reporte de incidencias."
)
OPERABILITY_NEW_2 = (
    "Para los alumnos de nuevo ingreso, la operacion debe considerar que todavia no conocen el "
    "proceso institucional de acceso ni el portal. El area administrativa debe cargar o vincular "
    "previamente la matricula y crear la cuenta en estado pendiente. El alumno solo debe confirmar "
    "su identidad, establecer su contrasena y activar la cuenta; no puede modificar matricula, rol "
    "ni permisos. Despues puede consultar su credencial QR, recuperar el acceso y recibir apoyo durante la induccion."
)
OPERABILITY_NEW_3 = (
    "La capacitacion minima debe cubrir el inicio de sesion, el uso del scanner, la interpretacion "
    "de resultados, la busqueda de registros, la creacion de Hot-QR, la gestion de personas, la "
    "revision de vehiculos, el cierre automatico, la revocacion de credenciales y el reporte de "
    "incidencias. Para alumnos de nuevo ingreso se agregara una orientacion breve sobre activacion, "
    "uso del QR y canales de soporte."
)

SCOPE_PARAGRAPHS = [
    "El alcance de identidad y administracion comprende autenticar administradores y usuarios del portal mediante sesiones seguras; proteger las sesiones con cookies httpOnly; administrar personas, tipos de persona, carreras, fotografias de perfil y estados de usuario; y separar los permisos de super administrador y administrador.",
    "El alcance de credenciales comprende generar, rotar y revocar credenciales QR personales; generar credenciales temporales y vehiculares; gestionar visitantes mediante Hot-QR; controlar vigencia y limite de usos; y revocar credenciales comprometidas.",
    "El alcance operativo comprende registrar entradas, salidas y rechazos desde el scanner; procesar Hot-QR y QR vehiculares; permitir captura manual auditada; registrar eventos, asistencia, vehiculos y permisos vehiculares; y conservar integridad mediante transacciones atomicas y cadena hash en PostgreSQL.",
    "El alcance administrativo comprende mostrar, filtrar y paginar registros; auditar acciones administrativas; automatizar expiraciones, cierres, limpieza y mantenimiento mediante el worker; probar flujos criticos con bun:test; y mostrar un contador de codigos QR escaneados con filtros por periodo, resultado y dispositivo.",
]

LIMITATION_PARAGRAPHS = [
    "La primera limitacion es temporal: la planeacion academica se realiza en cuatro semanas, mientras que la construccion completa se programa en dieciseis semanas. Cualquier ampliacion de alcance debe pasar por control de cambios para evitar desplazar pruebas, documentacion o entrega.",
    "La segunda limitacion es economica y tecnica: el proyecto depende de presupuesto autorizado, equipos prestados, servicios de infraestructura y disponibilidad de dispositivos compatibles para scanner. La precision de reportes historicos depende de la calidad de los datos capturados o migrados.",
    "La tercera limitacion es operativa: el despliegue requiere conectividad, servidor y base de datos disponibles, ademas de capacitacion para guardias, administradores y alumnos de nuevo ingreso. Las barreras automatizadas, torniquetes y otras integraciones fisicas quedan fuera de esta etapa.",
]

STATE_ART_HEADERS = ["Solucion", "Funciones principales", "Ventaja", "Limitacion", "Modalidad", "Aporte para v2"]
STATE_ART_ROWS = [
    ["Kisi", "Acceso en nube, app movil, hardware, permisos y reportes.", "Administracion remota y escalable.", "Dependencia de hardware y servicio externo.", "Paga: suscripcion y hardware bajo cotizacion.", "Permisos claros con datos bajo control institucional."],
    ["HID", "Credenciales fisicas y moviles, biometria e identidad.", "Solidez industrial y experiencia en campus.", "Implementacion especializada y costo alto.", "Paga: licencias, integracion y hardware.", "Identidad confiable mediante QR firmado y portal web."],
    ["Verkada", "Acceso hibrido, video, alertas y dispositivos.", "Eventos vinculados con contexto visual.", "Plataforma cerrada y costo empresarial.", "Paga: hardware y licencia recurrente.", "Auditoria y evidencia sin exponer datos innecesarios."],
    ["Avigilon/Openpath", "Credenciales moviles, pases, lectores e integraciones.", "Arquitectura completa de acceso fisico.", "Infraestructura profesional fuera del alcance inicial.", "Paga: suscripcion, lectores e instalacion.", "Modelo preparado para puertas futuras sin exigirlas ahora."],
]

SCHEDULE_HEADERS = ["Sem.", "Fechas", "Trabajo y avance", "Responsables", "Revision y evidencia"]
SCHEDULE_ROWS = [
    ["1", "03-09 ago 2026", "Inicio, problema, actores, objetivos y alcance preliminar.", "Lider + Analista", "Revision de propuesta y minuta de acuerdos."],
    ["2", "10-16 ago 2026", "Antecedentes, estado del arte, variables y delimitacion.", "Analista + QA", "Revision documental y correcciones de asesor."],
    ["3", "17-23 ago 2026", "Arquitectura, datos, latencia, usuarios, plataforma y presupuesto.", "Arquitecto + Frontend", "Revision de diseno y decisiones tecnicas."],
    ["4", "24-30 ago 2026", "Factibilidades, recursos, riesgos, PDCA y documento de planeacion.", "Todo el equipo", "Cierre de planeacion, evidencia y version aprobada."],
    ["5", "31 ago-06 sep 2026", "Monorepo, ambientes, esquema de datos y migraciones iniciales.", "Arquitecto + Lider", "Demo tecnica y revision de migraciones."],
    ["6", "07-13 sep 2026", "Autenticacion, sesiones, super admin, admin y permisos.", "Backend + QA", "Pruebas de seguridad y revision de roles."],
    ["7", "14-20 sep 2026", "Personas, matriculas, activacion de cuenta y catalogos.", "Analista + Backend", "Revision funcional con casos de nuevo ingreso."],
    ["8", "21-27 sep 2026", "QR personal, temporal, vehicular, expiracion y jti unico.", "Backend + QA", "Pruebas de firma, expiracion y consumo atomico."],
    ["9", "28 sep-04 oct 2026", "Scanner: entradas, salidas, rechazos y captura manual auditada.", "Frontend + Backend", "Demostracion del flujo completo de acceso."],
    ["10", "05-11 oct 2026", "Panel administrativo, tablas, filtros, paginacion y auditoria.", "Frontend + QA", "Revision de usabilidad y permisos."],
    ["11", "12-18 oct 2026", "Hot-QR, visitantes, vehiculos y permisos vehiculares.", "Backend + Frontend", "Pruebas funcionales y correccion de incidencias."],
    ["12", "19-25 oct 2026", "Worker, cierres automaticos, limpieza, mantenimiento y respaldos.", "Arquitecto + QA", "Revision de jobs, logs y recuperacion."],
    ["13", "26 oct-01 nov 2026", "Contador de escaneos, reportes y analitica operativa.", "Frontend + Analista", "Validacion de conteos por periodo y resultado."],
    ["14", "02-08 nov 2026", "Pruebas integrales, rendimiento, seguridad y accesibilidad.", "QA + Todo el equipo", "Reporte de pruebas e incidencias priorizadas."],
    ["15", "09-15 nov 2026", "Correcciones, capacitacion, manuales y revision completa del documento.", "Todo el equipo", "Ensayo operativo y revision APS final."],
    ["16", "16-22 nov 2026", "Despliegue, aceptacion, evidencias, presentacion y cierre.", "Lider + Todo el equipo", "Acta de cierre, entrega final y retrospectiva."],
]

TRACKING_ROWS = [[row[0], row[2], "", "", ""] for row in SCHEDULE_ROWS]

TOC_PAGES = {
    "1. Integrantes y roles": "4",
    "2. Descripcion general del proyecto": "4",
    "3. Planteamiento del problema": "5",
    "4. Objetivos": "5",
    "5. Antecedentes": "6",
    "6. Estado del arte": "7",
    "7. Justificacion": "8",
    "8. Variables del proyecto": "8",
    "9. Delimitacion del proyecto": "9",
    "10. Alcances": "9",
    "11. Limitaciones": "10",
    "12. Planeacion inicial": "11",
    "13. Planeacion economica y estimacion de costos": "13",
    "14. Factibilidad tecnica": "14",
    "15. Factibilidad operativa": "15",
    "16. Recursos del proyecto": "16",
    "17. Riesgos iniciales": "16",
    "18. Direccion, comunicacion y control del proyecto": "17",
    "19. Anexos": "17",
    "20. Conclusion": "18",
    "21. Bibliografia": "19",
}

CONCLUSION_OLD_2 = (
    "El principal reto es equilibrar seguridad, rapidez operativa y costo. El sistema debe ser "
    "suficientemente robusto para proteger credenciales, sesiones y auditoria, pero lo bastante "
    "simple para que guardias y administradores lo adopten sin friccion. Desde la factibilidad "
    "tecnica, operativa y economica, el proyecto es viable si se respeta el alcance inicial, se "
    "priorizan los flujos criticos y se deja la integracion fisica avanzada para fases posteriores."
)

ECONOMIC_HEADERS = [
    "Concepto", "Unidad de cobro", "Precio unitario", "Periodicidad", "Costo del proyecto", "Justificacion",
]
ECONOMIC_ROWS = [
    ["Lider del proyecto", "1 persona", "$40,000 al mes", "Mensual x 4", "$160,000", "Direccion, alcance, seguimiento y cierre."],
    ["Analista de requerimientos", "1 persona", "$35,000 al mes", "Mensual x 4", "$140,000", "Procesos, requisitos y aceptacion."],
    ["Arquitecto backend y datos", "1 persona", "$50,000 al mes", "Mensual x 4", "$200,000", "Arquitectura, seguridad, API y PostgreSQL."],
    ["Desarrollador frontend", "1 persona", "$35,000 al mes", "Mensual x 4", "$140,000", "Panel, scanner, portal y experiencia movil."],
    ["QA y documentacion", "1 persona", "$30,000 al mes", "Mensual x 4", "$120,000", "Pruebas, incidencias, evidencia y documentos."],
    ["Subtotal de salarios", "Equipo de 5 personas", "$190,000 al mes", "Mensual x 4", "$760,000", "Trabajo profesional directo."],
    ["Cargas y prestaciones estimadas", "30% de salarios", "$57,000 al mes", "Mensual x 4", "$228,000", "Costo patronal, administracion y reserva laboral."],
    ["Adquisicion y operacion", "1 proyecto", "$64,531.03", "Proyecto + 12 meses", "$64,531.03", "Hardware, servicios, mantenimiento y contingencia."],
    ["Trabajo en casa/hibrido", "5 personas", "$2,500 al mes equipo", "Mensual x 4", "$10,000", "Electricidad y consumibles; sin renta de oficina."],
    ["Proyecto base remoto/hibrido", "Proyecto completo", "$265,632.76 promedio", "4 meses", "$1,062,531.03", "Incluye salarios, cargas, adquisicion y trabajo en casa."],
    ["Oficina alternativa", "Espacio para 5 personas", "$15,000 al mes", "Mensual x 4", "$60,000", "Sustituye el costo de trabajo en casa; renta y servicios incluidos."],
    ["Proyecto maximo con oficina", "Proyecto completo", "$278,132.76 promedio", "4 meses", "$1,112,531.03", "Escenario superior para autorizacion presupuestal."],
]

RESOURCE_HEADERS = [
    "Recurso", "Modalidad", "Unidad de cobro", "Precio", "Periodicidad", "Costo del proyecto",
]
RESOURCE_ROWS = [
    ["Figma Education", "Gratuito referencial", "1 persona", "$0.01", "Mensual x 12", "$0.60"],
    ["GitHub Education", "Gratuito referencial", "1 persona", "$0.01", "Mensual x 12", "$0.60"],
    ["Visual Studio Code", "Gratuito referencial", "1 persona", "$0.01", "Mensual x 12", "$0.60"],
    ["Bun, TypeScript, SvelteKit, Hono", "Codigo abierto", "1 proyecto", "$0.01", "Mensual x 12", "$0.12"],
    ["Drizzle, Zod, Tailwind, librerias QR", "Codigo abierto", "1 proyecto", "$0.01", "Mensual x 12", "$0.12"],
    ["GitHub Projects", "Gratuito referencial", "1 persona", "$0.01", "Mensual x 12", "$0.60"],
    ["Ofimatica o documentacion", "Licencia existente", "1 persona", "$0.01", "Mensual x 12", "$0.60"],
    ["Servidor/VPS", "Servicio de pago", "1 proyecto", "$800", "Mensual x 12", "$9,600"],
    ["PostgreSQL administrado", "Servicio de pago", "1 proyecto", "$600", "Mensual x 12", "$7,200"],
    ["Dominio", "Servicio de pago", "1 proyecto", "$800", "Anual", "$800"],
    ["Storage R2/S3", "Servicio de pago", "1 proyecto", "$150", "Mensual x 12", "$1,800"],
    ["Scanner movil", "Equipo de pago", "2 dispositivos", "$4,500", "Una vez", "$9,000"],
    ["Capacitacion", "Servicio de pago", "10 h para equipo", "$180", "Una vez", "$1,800"],
    ["Modalidad de trabajo", "Remota/hibrida", "5 personas", "$2,500 equipo", "Mensual x 4", "$10,000"],
    ["Oficina alternativa", "Servicio de pago opcional", "1 espacio", "$15,000", "Mensual x 4", "$60,000"],
]


def node_text(node: etree._Element) -> str:
    return "".join(node.xpath(".//w:t/text()", namespaces=NS))


def replace_text(node: etree._Element, text: str) -> None:
    texts = node.xpath(".//w:t", namespaces=NS)
    if not texts:
        run = etree.SubElement(node, W + "r")
        target = etree.SubElement(run, W + "t")
        target.text = text
        return
    texts[0].text = text
    for target in texts[1:]:
        target.text = ""


def _set_table_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = DocxElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", 70), ("start", 90), ("bottom", 70), ("end", 90)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = DocxElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def build_table_xml(headers: list[str], rows: list[list[str]], widths: list[float], font_size=7.2) -> etree._Element:
    helper = Document()
    table = helper.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    table.rows[0]._tr.get_or_add_trPr().append(DocxElement("w:tblHeader"))
    for row_index, row in enumerate(table.rows):
        row._tr.get_or_add_trPr().append(DocxElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = DocxElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            _set_table_cell_margins(cell)
            if row_index == 0:
                shading = DocxElement("w:shd")
                shading.set(qn("w:fill"), "E8EEF5")
                tc_pr.append(shading)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = DocxElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    return etree.fromstring(etree.tostring(table._tbl))


def replace_table(root: etree._Element, marker: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    matches = [table for table in root.xpath(".//w:tbl", namespaces=NS) if marker in node_text(table)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table for {marker!r}; found {len(matches)}")
    old_table = matches[0]
    old_table.getparent().replace(old_table, build_table_xml(headers, rows, widths))


def find_paragraph(root: etree._Element, exact_text: str, *, direct_body: bool = False) -> etree._Element:
    path = "./w:body/w:p" if direct_body else ".//w:p"
    matches = [p for p in root.xpath(path, namespaces=NS) if node_text(p) == exact_text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph for: {exact_text[:80]!r}; found {len(matches)}")
    return matches[0]


def replace_section_with_paragraphs(
    root: etree._Element,
    heading_text: str,
    next_heading_text: str,
    paragraphs: list[str],
) -> None:
    body = root.find("w:body", NS)
    if body is None:
        raise RuntimeError("Missing document body")
    heading_scope = find_paragraph(root, heading_text, direct_body=True)
    heading_limitations = find_paragraph(root, next_heading_text, direct_body=True)
    children = list(body)
    start = children.index(heading_scope) + 1
    end = children.index(heading_limitations)
    scope_nodes = [node for node in children[start:end] if node.tag == W + "p" and node_text(node).strip()]
    if not scope_nodes:
        raise RuntimeError("No scope paragraphs found")
    template = deepcopy(find_paragraph(root, CONCLUSION_OLD_2, direct_body=True))
    for node in scope_nodes:
        body.remove(node)
    insertion_index = list(body).index(heading_limitations)
    for item in paragraphs:
        paragraph = deepcopy(template)
        replace_text(paragraph, item)
        body.insert(insertion_index, paragraph)
        insertion_index += 1


def update_state_of_art(root: etree._Element) -> None:
    replace_table(root, "Kisi", STATE_ART_HEADERS, STATE_ART_ROWS, [0.65, 1.25, 1.0, 1.15, 1.15, 1.3])


def update_schedule(root: etree._Element) -> None:
    replace_table(root, "Presentación del proyecto", SCHEDULE_HEADERS, SCHEDULE_ROWS, [0.45, 0.9, 2.25, 1.15, 1.75])
    replace_table(root, "Problema, objetivos, antecedentes", ["Semana", "Avance esperado", "Cumplio", "Observaciones", "Firma"], TRACKING_ROWS, [0.55, 3.0, 0.7, 1.45, 0.8])


def update_schedule_text(root: etree._Element) -> None:
    replacements = {
        "Tabla 5. Cronograma inicial de ocho semanas": "Tabla 5. Cronograma completo de dieciseis semanas",
        "Tabla 5. Cronograma inicial de cuatro semanas": "Tabla 5. Cronograma completo de dieciseis semanas",
        "3. Establecer un cronograma inicial de cuatro semanas con actividades, responsables,": "3. Establecer un cronograma completo de dieciseis semanas con fechas, actividades, responsables,",
        "Seguimiento del tiempo: cronograma de cuatro semanas y revision de lo planeado": "Seguimiento del tiempo: cronograma de dieciseis semanas y revision de lo planeado",
    }
    for paragraph in root.xpath(".//w:p", namespaces=NS):
        text = node_text(paragraph)
        for old, new in replacements.items():
            if text.startswith(old):
                replace_text(paragraph, text.replace(old, new, 1))
                break
    intro = next(
        p for p in root.xpath("./w:body/w:p", namespaces=NS)
        if node_text(p).startswith("La planeacion inicial se construye con base en los objetivos")
    )
    replace_text(
        intro,
        "La planeacion inicial se construye con base en objetivos, actividades, recursos, tiempo, responsables, comunicacion y control. Las primeras cuatro semanas corresponden a la planeacion academica; el cronograma completo cubre dieciseis semanas e integra diseno, desarrollo, pruebas, revisiones documentales, avances, capacitacion, despliegue y cierre.",
    )


def add_role_definition(root: etree._Element) -> None:
    marker = (
        "El proyecto es un sistema de informacion administrativo y operativo. Su proposito "
        "principal es reducir el uso de registros manuales, mejorar la seguridad del campus, "
        "conservar evidencia auditable de cada acceso y ofrecer una experiencia rapida para el "
        "punto fisico de entrada."
    )
    paragraph = find_paragraph(root, marker, direct_body=True)
    extra = deepcopy(paragraph)
    replace_text(
        extra,
        "El super administrador gobierna la configuracion global, crea o revoca administradores, asigna permisos, consulta toda la auditoria y atiende recuperaciones criticas. El administrador opera personas, matriculas, visitantes, vehiculos, accesos y reportes dentro de los permisos asignados; no puede crear super administradores, cambiar controles globales ni borrar evidencia de auditoria.",
    )
    parent = paragraph.getparent()
    parent.insert(parent.index(paragraph) + 1, extra)


def append_risk(root: etree._Element) -> None:
    tables = [table for table in root.xpath(".//w:tbl", namespaces=NS) if "QR reutilizado" in node_text(table)]
    if len(tables) != 1:
        raise RuntimeError(f"Risk table not found: {len(tables)}")
    table = tables[0]
    template = deepcopy(table.xpath("./w:tr", namespaces=NS)[-1])
    cells = template.xpath("./w:tc", namespaces=NS)
    values = [
        "Incapacidades o renuncias no planeadas",
        "Media",
        "Alto",
        "Mantener documentacion, respaldo de tareas, suplente por funcion critica, recuperacion de equipos y accesos, rotacion de secretos y replaneacion inmediata.",
    ]
    for cell, value in zip(cells, values):
        replace_text(cell, value)
    table.append(template)


def update_annexes(root: etree._Element) -> None:
    source = next(p for p in root.xpath(".//w:p", namespaces=NS) if node_text(p).startswith("Comunicacion:"))
    source_ppr = source.find("w:pPr", NS)
    for old, new in (
        ("Anexo A. Fuentes del curso analizadas:", "Fuentes del curso analizadas:"),
        ("Anexo B. Fuentes del proyecto analizadas:", "Fuentes del proyecto analizadas:"),
    ):
        paragraph = next(p for p in root.xpath("./w:body/w:p", namespaces=NS) if node_text(p).startswith(old))
        replace_text(paragraph, node_text(paragraph).replace(old, new, 1))
        current_ppr = paragraph.find("w:pPr", NS)
        if current_ppr is not None:
            paragraph.remove(current_ppr)
        if source_ppr is not None:
            paragraph.insert(0, deepcopy(source_ppr))


def ensure_table_vertical_borders(root: etree._Element) -> None:
    for table in root.xpath(".//w:tbl", namespaces=NS):
        tbl_pr = table.find("w:tblPr", NS)
        if tbl_pr is None:
            tbl_pr = etree.SubElement(table, W + "tblPr")
        borders = tbl_pr.find("w:tblBorders", NS)
        if borders is None:
            borders = etree.SubElement(tbl_pr, W + "tblBorders")
        for name in ("left", "right", "insideV"):
            border = borders.find(f"w:{name}", NS)
            if border is None:
                border = etree.SubElement(borders, W + name)
            border.set(W + "val", "single")
            border.set(W + "sz", "6")
            border.set(W + "space", "0")
            border.set(W + "color", "A6B7C8")


def remove_break_before_conclusion(root: etree._Element) -> None:
    heading = find_paragraph(root, "20. Conclusion", direct_body=True)
    previous = heading.getprevious()
    if previous is not None and previous.tag == W + "p" and previous.xpath(".//w:br[@w:type='page']", namespaces=NS):
        previous.getparent().remove(previous)


def add_qr_counter_variable(root: etree._Element) -> None:
    tables = root.xpath(".//w:tbl", namespaces=NS)
    variable_table = None
    for table in tables:
        text = node_text(table)
        if "Variable" in text and "Matricula/identificador" in text and "Configuracion" in text:
            variable_table = table
            break
    if variable_table is None:
        raise RuntimeError("Variable table not found")
    rows = variable_table.xpath("./w:tr", namespaces=NS)
    template = deepcopy(rows[-1])
    cells = template.xpath("./w:tc", namespaces=NS)
    values = [
        "Contador de escaneos QR",
        "Total calculado a partir de eventos de scanner; permite consulta por periodo, resultado o dispositivo.",
        "Analitica operativa",
    ]
    if len(cells) != len(values):
        raise RuntimeError("Unexpected variable table geometry")
    for cell, value in zip(cells, values):
        replace_text(cell, value)
    variable_table.append(template)


def add_operability_paragraph(root: etree._Element) -> None:
    first = find_paragraph(root, OPERABILITY_OLD_1)
    second = find_paragraph(root, OPERABILITY_OLD_2)
    replace_text(first, OPERABILITY_NEW_1)
    replace_text(second, OPERABILITY_NEW_2)
    parent = second.getparent()
    extra = deepcopy(second)
    replace_text(extra, OPERABILITY_NEW_3)
    parent.insert(parent.index(second) + 1, extra)


def update_financial_tables(root: etree._Element) -> None:
    replace_table(root, "Horas de desarrollo", ECONOMIC_HEADERS, ECONOMIC_ROWS, [1.1, 1.15, 1.0, 0.8, 1.05, 1.4])
    replace_table(root, "Recursos gratuitos", RESOURCE_HEADERS, RESOURCE_ROWS, [1.3, 1.0, 1.0, 0.7, 1.1, 1.4])


def extend_conclusion(root: etree._Element) -> None:
    second = find_paragraph(root, CONCLUSION_OLD_2, direct_body=True)
    parent = second.getparent()
    additions = [
        "La planeacion economica debe leerse por unidad de cobro para evitar mezclar costos personales con costos del proyecto. El trabajo profesional se estima por rol y por mes durante cuatro meses; los equipos prestados se valoran al 80%; la infraestructura se cobra por proyecto y por periodo; y las compras fisicas se registran una sola vez. Esta separacion permite identificar costos mensuales, anuales, por persona, por equipo y por proyecto completo.",
        "El presupuesto tambien distingue el desembolso efectivo de la aportacion en especie. Las herramientas gratuitas se registran con un valor referencial de $0.01 MXN para conservar trazabilidad sin tratarlas como una compra. En cada revision semanal el equipo debe comparar horas consumidas, servicios activos, dispositivos disponibles y contingencia restante contra el presupuesto aprobado; cualquier cambio debe documentarse antes de comprometer gasto adicional.",
        "Con estos controles, Control Acceso UPQROO v2 puede completar cuatro semanas de planeacion y dieciseis semanas de ejecucion con responsables, revisiones y evidencias definidas. El costo base remoto o hibrido se estima en $1,062,531.03 MXN y el escenario maximo con oficina en $1,112,531.03 MXN. La entrega debe conservar esta disciplina al contratar servicios, mantener equipos prestados, medir escaneos QR y capacitar a los usuarios, en especial a los alumnos de nuevo ingreso.",
    ]
    insert_at = parent.index(second) + 1
    for text in additions:
        paragraph = deepcopy(second)
        replace_text(paragraph, text)
        parent.insert(insert_at, paragraph)
        insert_at += 1


def update_toc_pages(root: etree._Element) -> None:
    toc = None
    for table in root.xpath(".//w:tbl", namespaces=NS):
        if "Seccion" in node_text(table) and "Pagina" in node_text(table) and "1. Integrantes y roles" in node_text(table):
            toc = table
            break
    if toc is None:
        raise RuntimeError("Table of contents not found")
    for row in toc.xpath("./w:tr", namespaces=NS)[1:]:
        cells = row.xpath("./w:tc", namespaces=NS)
        if len(cells) != 2:
            continue
        section = node_text(cells[0])
        if section in TOC_PAGES:
            replace_text(cells[1], TOC_PAGES[section])


def patch_document_xml(xml_bytes: bytes) -> bytes:
    parser = etree.XMLParser(remove_blank_text=False)
    root = etree.fromstring(xml_bytes, parser)
    replace_text(find_paragraph(root, OBJECTIVE_OLD), OBJECTIVE_NEW)
    add_role_definition(root)
    add_qr_counter_variable(root)
    update_state_of_art(root)
    replace_section_with_paragraphs(root, "10. Alcances", "11. Limitaciones", SCOPE_PARAGRAPHS)
    replace_section_with_paragraphs(root, "11. Limitaciones", "12. Planeacion inicial", LIMITATION_PARAGRAPHS)
    update_schedule(root)
    update_schedule_text(root)
    add_operability_paragraph(root)
    update_financial_tables(root)
    append_risk(root)
    update_annexes(root)
    extend_conclusion(root)
    remove_break_before_conclusion(root)
    update_toc_pages(root)
    ensure_table_vertical_borders(root)
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")


def write_corrected_docx() -> None:
    with ZipFile(SOURCE, "r") as source_zip, ZipFile(CORRECTED, "w") as output_zip:
        for info in source_zip.infolist():
            data = source_zip.read(info.filename)
            if info.filename == "word/document.xml":
                data = patch_document_xml(data)
            new_info = ZipInfo(info.filename, date_time=info.date_time)
            new_info.compress_type = info.compress_type or ZIP_DEFLATED
            new_info.comment = info.comment
            new_info.extra = info.extra
            new_info.internal_attr = info.internal_attr
            new_info.external_attr = info.external_attr
            new_info.create_system = info.create_system
            output_zip.writestr(new_info, data)


if __name__ == "__main__":
    write_corrected_docx()
    print(CORRECTED)
