from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("plan_adquisicion_equipos_y_costos_control_acceso_v2.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"


def set_font(run, size=None, bold=None, color=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    total = int(sum(widths) * 1440)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "A6B7C8")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[index] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Control Acceso UPQROO v2 | Adquisicion de equipos y costos")
    set_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Pagina ")
    set_font(run, size=9, color=MUTED)
    add_field(footer, "PAGE")
    run = footer.add_run(" de ")
    set_font(run, size=9, color=MUTED)
    add_field(footer, "NUMPAGES")
    for footer_run in footer.runs:
        set_font(footer_run, size=9, color=MUTED)


def add_title(doc: Document, text: str, size: float, after: float, color=BLACK) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=True, color=color)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, size=11, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_font(run, size=11)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(f"{label}: ")
    set_font(run, size=11, bold=True, color=DARK_BLUE)
    run = paragraph.add_run(text)
    set_font(run, size=11)
    set_table_geometry(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size=8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, HEADER_FILL)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=font_size, bold=(row_index == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=10, bold=True, color=DARK_BLUE)
    run.italic = True


def build() -> None:
    doc = Document()
    style_document(doc)

    add_title(doc, "UNIVERSIDAD POLITECNICA DE QUINTANA ROO", 15, 20)
    add_title(doc, "Administracion de Proyectos de Software", 14, 10)
    add_title(doc, "Plan de Adquisicion de Equipos y Costos", 24, 6, DARK_BLUE)
    add_title(doc, "Sistema de Control de Acceso UPQROO v2", 18, 28)
    for label, value in (
        ("Equipo", "Equipo Control Acceso v2"),
        ("Periodo de planeacion", "Cuatro semanas"),
        ("Horizonte de costos", "Implementacion inicial y operacion por 12 meses"),
        ("Moneda", "Pesos mexicanos (MXN)"),
        ("Fecha", "Julio de 2026"),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(f"{label}: ")
        set_font(run, size=11, bold=True)
        run = paragraph.add_run(value)
        set_font(run, size=11)
    doc.add_page_break()

    doc.add_heading("1. Proposito y criterio de costeo", level=1)
    add_body(doc, "Este documento concentra los equipos, accesorios, servicios, herramientas, conectividad, licencias y costos necesarios para desarrollar, probar, desplegar y operar Control Acceso UPQROO v2. Distingue entre compra, prestamo, recurso ya disponible y servicio recurrente para evitar contabilizar como gasto en efectivo aquello que el equipo aporta temporalmente.")
    add_callout(doc, "Criterio de decision", "Se prioriza el prestamo documentado de equipos propios y los planes educativos o gratuitos. La compra se reserva para elementos que afectan continuidad, seguridad, respaldo, cableado o la operacion del scanner.")
    add_body(doc, "Los importes son estimaciones academicas de planeacion. Antes de comprar se deben solicitar cotizaciones vigentes, confirmar impuestos, validar compatibilidad y registrar responsable, numero de serie, condicion de entrega y fecha de devolucion.")

    doc.add_heading("2. Supuestos de disponibilidad", level=1)
    add_caption(doc, "Tabla 1. Recursos disponibles y tratamiento presupuestal")
    add_table(doc, ["Recurso", "Unidad", "Periodicidad", "Modalidad", "Costo efectivo", "Tratamiento"], [
        ["Laptops de desarrollo", "5 personas", "Durante el proyecto", "Prestamo", "$14,400/equipo", "Valor de uso equivalente al 80% de $18,000; registrar serie, cargador y estado fisico."],
        ["Telefonos para scanner", "2 dispositivos", "Durante el proyecto", "Prestamo", "$3,600/equipo", "Valor de uso equivalente al 80% de $4,500; comprobar camara, bateria y navegador."],
        ["Correo institucional", "5 personas", "Mensual x 12", "Disponible", "$0.01/persona", "Usar para Figma Education, GitHub Education y comunicacion."],
    ], [1.15, 0.75, 1.0, 0.8, 0.95, 1.85], font_size=7.5)
    add_body(doc, "Los equipos prestados se registran al 80% de su precio de referencia para representar su uso temporal y desgaste sin tratarlos como compra nueva. El valor de $0.01 MXN queda reservado exclusivamente para recursos gratuitos o licencias disponibles.")

    doc.add_heading("3. Equipos aportados en prestamo", level=1)
    add_caption(doc, "Tabla 2. Inventario previsto de prestamo")
    add_table(doc, ["Bien", "Cantidad", "Precio de referencia", "Valor al 80%", "Registro requerido"], [
        ["Laptop personal de desarrollo", "5", "$18,000", "$72,000", "Responsable, marca, modelo, serie, cargador, estado y fechas."],
        ["Cargador original de laptop", "5", "$1,200", "$4,800", "Asociar al numero de inventario de la laptop."],
        ["Telefono o tablet para scanner", "2", "$4,500", "$7,200", "Camara funcional, navegador actualizado, bateria y responsable."],
        ["Telefono de respaldo", "1", "$4,500", "$3,600", "Disponible durante pruebas y demostracion final."],
    ], [1.65, 0.65, 0.95, 0.95, 2.3])
    add_callout(doc, "Valor total aportado en prestamo", "$87,600 MXN, equivalente al 80% del precio de referencia de los equipos. Este valor reconoce uso y desgaste; no representa un desembolso si los bienes se devuelven conforme al acuerdo.")
    add_caption(doc, "Tabla 3. Control de responsables del prestamo")
    add_table(doc, ["Responsable", "Equipo asignado", "Serie/identificador", "Entrega", "Devolucion"], [
        ["Cruz Flores Juan Leonardo", "Laptop personal y cargador", "Por completar", "Por completar", "Por completar"],
        ["Gonzales Camara Didier Ariel", "Laptop personal y cargador", "Por completar", "Por completar", "Por completar"],
        ["Pech Canche Raul Gerardo", "Laptop personal y cargador", "Por completar", "Por completar", "Por completar"],
        ["Gongora Canul Raul Noe", "Laptop personal y cargador", "Por completar", "Por completar", "Por completar"],
        ["Prestegui Cruz Angel Gael", "Laptop personal y cargador", "Por completar", "Por completar", "Por completar"],
        ["Custodio por definir", "Dos scanners y un equipo de respaldo", "Por completar", "Por completar", "Por completar"],
    ], [1.6, 2.0, 1.15, 0.85, 0.9])

    doc.add_heading("4. Compras de hardware y accesorios", level=1)
    add_caption(doc, "Tabla 4. Adquisiciones recomendadas")
    add_table(doc, ["Concepto", "Unidad de cobro", "Precio unitario", "Periodicidad", "Costo proyecto", "Necesidad"], [
        ["Punto de acceso movil o hotspot", "1 dispositivo", "$1,200", "Una vez", "$1,200", "Contingencia cuando falle la red fija."],
        ["Soporte para telefono o tablet", "2 dispositivos", "$350", "Una vez", "$700", "Mantener estable la camara en el punto de acceso."],
        ["Bateria externa de 20,000 mAh", "2 dispositivos", "$700", "Una vez", "$1,400", "Continuidad del scanner ante bateria baja."],
        ["Regulador o supresor de picos", "2 unidades", "$450", "Una vez", "$900", "Proteccion electrica para carga y pruebas."],
        ["UPS de respaldo", "1 unidad", "$2,500", "Una vez", "$2,500", "Respaldo para router o equipo local critico."],
        ["Camara web USB de respaldo", "2 unidades", "$650", "Una vez", "$1,300", "Alternativa si la camara integrada no funciona."],
        ["Bobina de cable UTP Cat 6 para exterior", "Longitud total maxima de 100 m", "$1,500", "Una vez", "$1,500", "Cableado exterior para pruebas o instalacion; no exceder 100 m."],
        ["Kit de terminacion RJ45 y prueba", "1 lote", "$600", "Una vez", "$600", "Conectores, botas, ponchadora y verificador basico."],
        ["Adaptador USB-C multipuerto", "1 unidad", "$700", "Una vez", "$700", "Conectar red, camara o almacenamiento."],
        ["SSD externo de 1 TB", "1 unidad", "$1,600", "Una vez", "$1,600", "Respaldo cifrado de entregables y evidencias."],
        ["Credenciales e impresion inicial", "1 lote", "$1,000", "Una vez", "$1,000", "Pruebas de QR, gafetes y material operativo."],
        ["Fundas, micas y limpieza", "1 lote", "$800", "Una vez", "$800", "Proteccion basica de scanners prestados."],
        ["Subtotal de hardware", "Proyecto", "", "Una vez", "$14,200", "Compra inicial recomendada; $2,840 por integrante si se divide entre 5."],
    ], [1.35, 1.0, 0.85, 0.85, 1.0, 1.45], font_size=7.4)
    add_body(doc, "Si los telefonos prestados no cumplen con camara, autonomia o navegador, se activa el escenario alterno de compra de dos dispositivos scanner por $9,000 MXN adicionales.")

    doc.add_heading("5. Servicios, conectividad y operacion", level=1)
    add_caption(doc, "Tabla 5. Servicios, unidad de cobro y periodicidad")
    add_table(doc, ["Servicio", "Unidad de cobro", "Precio", "Periodicidad", "Costo proyecto", "Responsable"], [
        ["Servidor VPS o hosting", "1 proyecto", "$800", "Mensual x 12", "$9,600", "Lider tecnico"],
        ["PostgreSQL administrado", "1 proyecto", "$600", "Mensual x 12", "$7,200", "Backend y datos"],
        ["Dominio", "1 proyecto", "$800", "Anual", "$800", "Lider del proyecto"],
        ["Cloudflare R2 o storage", "1 proyecto", "$150", "Mensual x 12", "$1,800", "Backend y datos"],
        ["Respaldo adicional", "1 proyecto", "$50", "Mensual x 12", "$600", "QA y documentacion"],
        ["Datos moviles", "2 lineas", "$250", "Mensual x 3", "$1,500", "Custodio de scanners"],
        ["HTTPS/TLS", "1 proyecto", "$0.01", "Anual", "$0.01", "Lider tecnico"],
        ["Correo institucional", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Cada integrante"],
        ["Subtotal de servicios", "Proyecto", "", "12 meses", "$21,500.61", "Infraestructura, almacenamiento, datos moviles y dominio."],
    ], [1.35, 0.95, 0.75, 1.0, 1.1, 1.35], font_size=7.3)

    doc.add_heading("6. Herramientas y licencias", level=1)
    add_caption(doc, "Tabla 6. Herramientas, personas usuarias y costo referencial")
    add_table(doc, ["Herramienta", "Unidad de cobro", "Precio", "Periodicidad", "Costo proyecto", "Plan o evidencia"], [
        ["Figma / FigJam", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Education verificado; Starter como alternativa."],
        ["GitHub", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Free o Student Developer Pack."],
        ["Visual Studio Code", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Instalacion en laptops prestadas."],
        ["Bun, TypeScript, SvelteKit, Hono", "1 proyecto", "$0.01", "Mensual x 12", "$0.12", "Versiones fijadas en monorepo."],
        ["PostgreSQL, Drizzle ORM, Zod", "1 proyecto", "$0.01", "Mensual x 12", "$0.12", "Uso local mas servicio administrado."],
        ["html5-qrcode, qrcode, jose", "1 proyecto", "$0.01", "Mensual x 12", "$0.12", "Revisar licencias y versiones."],
        ["GitHub Projects", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Seguimiento semanal y bitacora."],
        ["Documentacion y ofimatica", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Licencia institucional o software disponible."],
        ["Gestor de contrasenas", "5 personas", "$0.01", "Mensual x 12", "$0.60", "Plan gratuito para secretos del proyecto."],
        ["Subtotal de herramientas", "Proyecto", "", "12 meses", "$3.96", "Valor referencial solicitado; no es una compra."],
    ], [1.3, 0.95, 0.75, 1.0, 1.05, 1.45], font_size=7.2)
    add_body(doc, "La condicion de Figma no debe afirmarse como activa hasta verificarla. El responsable de diseno debe anexar captura o correo de aprobacion de Figma Education; si no existe, se documentara el uso del plan Starter. El valor de $0.01 MXN se usa solo como referencia para recursos gratuitos.")

    doc.add_heading("7. Costos de preparacion y soporte", level=1)
    add_caption(doc, "Tabla 7. Preparacion, costo por equipo y referencia por persona")
    add_table(doc, ["Actividad", "Unidad de cobro", "Precio", "Periodicidad", "Costo proyecto", "Referencia por persona"], [
        ["Configuracion e inventario", "12 h del equipo", "$180/h", "Una vez", "$2,160", "$432 si se divide entre 5"],
        ["Capacitacion operativa", "10 h del equipo", "$180/h", "Una vez", "$1,800", "$360 si se divide entre 5"],
        ["Subtotal de preparacion", "Proyecto", "", "Una vez", "$3,960", "$792 por integrante"],
    ], [1.35, 1.0, 0.75, 0.9, 1.05, 1.45], font_size=7.5)

    add_caption(doc, "Tabla 8. Mantenimiento preventivo y reserva correctiva")
    add_table(doc, ["Equipo o servicio", "Unidad", "Precio unitario", "Periodicidad", "Costo anual", "Cobertura"], [
        ["Laptops de desarrollo", "5 equipos", "$800 por servicio", "Semestral", "$8,000", "Limpieza interna, revision termica, sistema y diagnostico."],
        ["Telefonos o tablets scanner", "3 equipos", "$500 por servicio", "Semestral", "$3,000", "Limpieza, bateria, camara, puertos y actualizaciones."],
        ["UPS, reguladores y red", "1 lote", "$1,500 por revision", "Semestral", "$3,000", "Baterias, contactos, cableado y conectividad."],
        ["Reserva de mantenimiento correctivo", "1 proyecto", "$5,000", "Anual", "$5,000", "Diagnostico, mano de obra menor y sustitucion de accesorios."],
        ["Subtotal de mantenimiento", "Proyecto", "", "12 meses", "$19,000", "Presupuesto preventivo y correctivo; refacciones mayores requieren autorizacion."],
    ], [1.35, 0.85, 0.95, 0.9, 1.0, 1.45], font_size=7.3)

    doc.add_heading("8. Resumen economico", level=1)
    add_caption(doc, "Tabla 9. Presupuesto consolidado por responsable de costo")
    add_table(doc, ["Categoria", "Por persona", "Por equipo", "Proyecto efectivo", "Aporte en especie", "Periodicidad"], [
        ["Hardware y accesorios", "$2,840 ref.", "$14,200", "$14,200", "$87,600", "Una vez"],
        ["Servicios e infraestructura", "$4,300.12 ref.", "$21,500.61", "$21,500.61", "$0", "12 meses"],
        ["Herramientas gratuitas", "$0.79 ref.", "$3.96", "$3.96", "$0", "12 meses"],
        ["Preparacion y capacitacion", "$792", "$3,960", "$3,960", "$0", "Una vez"],
        ["Mantenimiento", "$3,800 ref.", "$19,000", "$19,000", "$0", "12 meses"],
        ["Subtotal de adquisicion", "$11,732.91 ref.", "$58,664.57", "$58,664.57", "$87,600", "Antes de contingencia"],
        ["Contingencia 10%", "$1,173.29 ref.", "$5,866.46", "$5,866.46", "$0", "Una vez"],
        ["Adquisicion y operacion", "$12,906.21 ref.", "$64,531.03", "$64,531.03", "$87,600", "Proyecto + 12 meses"],
        ["Trabajo profesional", "$197,600", "$988,000", "$988,000", "$0", "4 meses, cargas incluidas"],
        ["Proyecto base remoto/hibrido", "$212,506.21 ref.", "$1,062,531.03", "$1,062,531.03", "$87,600", "Incluye $10,000 de operacion en casa"],
        ["Proyecto maximo con oficina", "$222,506.21 ref.", "$1,112,531.03", "$1,112,531.03", "$87,600", "Sustituye casa por oficina durante 4 meses"],
    ], [1.25, 0.85, 0.95, 1.15, 1.1, 1.2], font_size=7.1)
    add_callout(doc, "Como leer el resumen", "El costo por persona es una referencia si el equipo divide los gastos entre cinco integrantes. El costo por equipo muestra el total compartido. El costo efectivo del proyecto indica el desembolso. La aportacion en especie refleja el valor de prestamo sin convertirlo en compra.")
    add_callout(doc, "Escenario de compra de scanners", "Si no se prestan dos dispositivos compatibles, el desembolso aumenta $9,000 MXN mas $900 de contingencia. El total de adquisicion y operacion pasaria aproximadamente a $74,431.03 MXN; el maximo del proyecto con oficina seria $1,122,431.03 MXN.")

    doc.add_heading("9. Procedimiento de adquisicion y prestamo", level=1)
    for label, text in (
        ("Inventario", "Registrar responsable, marca, modelo, serie, accesorios, estado fisico y evidencia fotografica."),
        ("Compatibilidad", "Comprobar navegador actualizado, camara, bateria, conectividad, almacenamiento disponible y bloqueo de pantalla."),
        ("Autorizacion", "Obtener firma de prestamo antes de instalar herramientas o usar el equipo en demostraciones."),
        ("Compra", "Solicitar al menos dos cotizaciones cuando el importe lo justifique; comparar garantia, entrega e impuestos."),
        ("Configuracion", "Aplicar actualizaciones, cuentas separadas, cifrado, antivirus, copias de seguridad y acceso minimo necesario."),
        ("Devolucion", "Eliminar credenciales del proyecto, entregar respaldos, verificar estado y firmar conformidad."),
    ):
        add_body(doc, f"{label}: {text}", bold_prefix=f"{label}: ")

    doc.add_heading("10. Criterios de aceptacion", level=1)
    add_caption(doc, "Tabla 10. Verificacion antes de operar")
    add_table(doc, ["Elemento", "Criterio minimo", "Evidencia"], [
        ["Laptop", "Compila el monorepo, ejecuta pruebas y conserva al menos 20 GB libres.", "Captura de version, prueba y numero de inventario."],
        ["Scanner", "Lee QR con iluminacion normal, mantiene sesion y opera durante una jornada de prueba.", "Video o bitacora de prueba."],
        ["Contingencia", "Hotspot y captura manual auditada disponibles conforme al plan.", "Prueba controlada de falla."],
        ["Figma", "Plan Education verificado o Starter documentado.", "Captura de cuenta y equipo de trabajo."],
        ["Prestamo", "Serie, estado, firma y fecha de devolucion completos.", "Formato firmado."],
        ["Respaldo", "Copia cifrada recuperable y responsable asignado.", "Prueba de restauracion."],
    ], [1.25, 3.45, 1.8])

    doc.add_heading("11. Riesgos de adquisicion", level=1)
    add_caption(doc, "Tabla 11. Riesgos y respuestas")
    add_table(doc, ["Riesgo", "Respuesta prevista"], [
        ["Equipo prestado deja de estar disponible", "Activar responsable alterno o comprar scanner conforme al escenario de sustitucion."],
        ["Internet institucional falla", "Usar hotspot y datos moviles; conservar captura manual auditada solo como contingencia."],
        ["No se aprueba Figma Education", "Usar Figma Starter sin costo; limitar archivos y colaboradores a lo necesario."],
        ["Costo de servicio cambia", "Cotizar antes de contratar y usar la contingencia sin reducir seguridad ni respaldos."],
        ["Danio o perdida de equipo", "Definir custodio, funda, transporte, bloqueo y acta de entrega-recepcion."],
        ["Cuenta personal conserva accesos", "Usar cuentas del proyecto, rotar secretos y revocar sesiones al finalizar."],
        ["Incapacidad temporal de un integrante", "Mantener documentacion, respaldo de tareas y una persona suplente por responsabilidad critica."],
        ["Renuncia no planeada", "Recuperar equipos y accesos, rotar secretos, reasignar entregables y actualizar el cronograma y presupuesto."],
    ], [2.25, 4.25])

    doc.add_heading("12. Conclusion", level=1)
    add_body(doc, "La estrategia recomendada combina prestamo documentado de equipos propios, herramientas gratuitas o educativas, compras puntuales, mantenimiento preventivo y una reserva correctiva. Los equipos prestados se reconocen por $87,600 MXN, equivalente al 80% de su precio de referencia, mientras que la adquisicion y operacion requieren $64,531.03 MXN. El escenario alterno de compra de scanners se activa solo si los dispositivos prestados no cumplen con los criterios tecnicos.")
    add_body(doc, "La separacion por persona, equipo de desarrollo y proyecto evita interpretar un costo anual como si fuera individual. El proyecto base remoto o hibrido se estima en $1,062,531.03 MXN y el escenario maximo con oficina en $1,112,531.03 MXN. Los servicios SaaS muestran unidad de cobro y periodicidad; las compras fisicas se registran una vez; los equipos prestados se valoran al 80%; los importes de $0.01 MXN se reservan para herramientas realmente gratuitas.")
    add_body(doc, "Antes de la entrega final deben completarse las series de los equipos, verificar el estado de Figma Education o documentar el plan Starter, obtener cotizaciones, programar mantenimiento y firmar los formatos de prestamo. En cada revision semanal se debe comparar el gasto real contra los servicios activos, los accesorios adquiridos, la contingencia disponible y la responsabilidad asignada a cada integrante.")

    doc.add_heading("13. Fuentes de referencia", level=1)
    for source in (
        "Figma, Figma for Education, acceso gratuito para estudiantes y educadores con estado educativo verificado: https://help.figma.com/hc/en-us/articles/360041061214-Figma-for-Education",
        "GitHub Docs, Student Developer Pack para estudiantes verificados: https://docs.github.com/en/education/about-github-education/github-education-for-students/github-terms-and-conditions-for-the-student-developer-pack",
        "Cloudflare Docs, R2 Pricing, nivel gratuito y precios por almacenamiento y operaciones: https://developers.cloudflare.com/r2/pricing/",
        "Zegucom Computo, bobina Cat 6 para exterior de 100 metros, precio de referencia consultado en julio de 2026: https://www.zegucom.com.mx/producto/cables-y-accesorios/cables-de-red-en-bobina/bobina-cable-de-red-x-case-c10021dofo-100m-cat6-u-utp-negro/QzEwMDIxRE9GTw%3D%3D",
        "Indeed Mexico, referencias salariales por rol de proyecto, arquitectura, frontend, analisis y QA, consultadas en julio de 2026: https://mx.indeed.com/career/qa-engineer/salaries",
        "Fase de Planeacion del Proyecto, Sistema de Control de Acceso UPQROO v2, julio de 2026.",
    ):
        add_body(doc, source)

    doc.core_properties.title = "Plan de adquisicion de equipos y costos - Control Acceso v2"
    doc.core_properties.subject = "Administracion de Proyectos de Software"
    doc.core_properties.author = "Equipo Control Acceso v2"
    doc.core_properties.keywords = "UPQROO, adquisicion, equipos, costos, prestamo, Control Acceso v2"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
